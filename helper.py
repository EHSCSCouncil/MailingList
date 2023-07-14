import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders
from openpyxl import load_workbook
from jinja2 import Template
import pandas as pd
from docx import Document
from dotenv import dotenv_values
def read_excel_data(excel_path):
    data = pd.read_excel(excel_path)
    columns = data.columns.tolist()
    to, cc, attachments, keywords = None, None, None, None
    if 'To' in columns:
        to = data['To']
    else:
        raise Exception("Need a column of 'To' in excel path")
    if 'CC' in columns:
        cc = data['CC']
    if 'Attachments' in columns:
        attachments = data[columns]
    keywords = {col: data[col] for col in columns if col not in ['To', 'CC', 'Attachments']}
    return to, cc, attachments, keywords

def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None

def read_word_data(word_path):
    doc = Document(word_path)
    first_line = doc.paragraphs[0]
    if 'Subject: ' not in first_line.text:
        raise Exception("First line of word file must start with 'Subject:'")
    subject = first_line.text.split("Subject: ")[1]
    delete_paragraph(first_line)
    text = '\n\n'.join([paragraph.text for paragraph in doc.paragraphs])
    print(text)
    return subject, text

def email_credentials():
    secrets = dotenv_values(".env")
    username = secrets["username"]
    password = secrets["password"]
    smtp_server = secrets["server"]
    smtp_port = secrets["port"]
    return username, password, smtp_server, smtp_port

def substitute_data(body, keywords):
    for key, value in keywords.items():
        body = body.replace('{{' + key + '}}', value)
    if '{{' in body or '}}' in body:
        raise Exception("there are leftover keywords")
    return body

def send_email(smtp_server, smtp_port, sender_email, sender_password, to, cc, subject, body, attachments=None):
    msg = MIMEMultipart()
    msg['From'] = sender_email
    msg['To'] = to
    msg['CC'] = cc
    msg['Subject'] = subject

    msg.attach(MIMEText(body, 'plain'))

    if attachments:
        for attachment in attachments:
            part = MIMEBase('application', 'octet-stream')
            part.set_payload(open(attachment, 'rb').read())
            encoders.encode_base64(part)
            part.add_header('Content-Disposition', 'attachment', filename=attachment)
            msg.attach(part)

    with smtplib.SMTP(smtp_server, smtp_port) as server:
        server.starttls()
        server.login(sender_email, sender_password)
        server.send_message(msg)
