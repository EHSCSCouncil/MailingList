import sys
from dotenv import dotenv_values
import openpyxl
from docx import Document
from dotenv import load_dotenv

secrets = dotenv_values(".env")
username = secrets["username"]
password = secrets["password"]


def get_recipients(cell_value):
   
    return cell_value.split(", ")

def process_files(excel_path, word_path):

    wb = openpyxl.load_workbook(excel_path)
    sheet = wb.active

    to_col_idx = sheet[1].index('To')
    cc_col_idx = sheet[1].index('CC')
    attachments_col_idx = sheet[1].index('Attachments')

    for row in sheet.iter_rows(min_row = 2, values_only = True):
        to_emails = get_recipients(row[to_col_idx]) 
        cc_emails = get_recipients(row[cc_col_idx])  
        attachments = row[attachments_col_idx].split(", ")

        doc = Document(word_path)

        subject = doc.paragraphs[0].text.split("Subject: ")[1]


        for idx, cell_value in enumerate(sheet[1], start=2):
            keyword = "{{" + cell_value + "}}"
            value = str(row[idx])
            for paragraph in doc.paragraphs:
                if keyword in paragraph.text:
                    paragraph.text = paragraph.text.replace(keyword, value)

        print("To:", to_emails)
        print("CC:", cc_emails)
        print("Subject:", subject)
        print("Attachments:", attachments)
        print("")




def main():
   print(details_path)
   print(template_path)
   try:
      details_path = sys.argv[1]
      template_path = sys.argv[2]
   except IndexError:
      raise SystemExit("add 2 paths as arguments: path to an excel spreadsheet and path to a word doc template")
   
   process_files(details_path, template_path)

if __name__ == "__main__":
   main()
